import sys
import os
import matplotlib.pyplot as plt

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH


from plot_params import ePlot_type
import pandas as pd

import pdb
class OutputReport():

    def __init__(self,**kwargs):
        
        
        self.loc=kwargs.get("loc",r"C:\test.docx")
        self.debug_mode=kwargs.get("debug_mode",False)

        self.file=docx.Document()
        self.file.save(self.loc)


    def add_plot(self,plot):#just takes a single plot() object

        if not plot.generated:
            plot.generate_plot()
            

        plot.figure.savefig("temp.png",bbox_inches='tight')

        for section in self.file.sections:
            width=8.5-section.left_margin.inches-section.right_margin.inches

        self.file.add_picture("temp.png",width=docx.shared.Inches(width))
        os.remove("temp.png")

        if self.debug_mode:
            plot.figure.show()
            


    def add_table(self,data,**kwargs):
        merges = kwargs.get("merges", [])
        if isinstance(data,pd.DataFrame):

            data=[data.columns.values.tolist()]+ data.values.tolist()
            
        table=self.file.add_table(rows=len(data),cols=len(data[0]))

        for i,_ in enumerate(data):
            for j,_ in enumerate(data[i]):
                cell=table.cell(i,j)
                cell.text=str(data[i][j])

        for merge in merges:
            a=table.cell(merge[0][0],merge[0][1])
            b=table.cell(merge[1][0],merge[1][1])
            A=a.merge(b)
        table.style = 'LightShading-Accent1'

        return table


    def add_text(self,text,**kwargs):
        the_type=kwargs.get("type","body")#body, title, header
        default=WD_ALIGN_PARAGRAPH.LEFT
        if the_type=="title":
            default=WD_ALIGN_PARAGRAPH.CENTER

        alignment=kwargs.get("align",default)


        if the_type=="title":
            print("adding title")
            t=self.file.add_heading(text,level=0)
            
        elif the_type=="heading":
            
            t=self.file.add_heading(text,level=1)
        else:#body
            
            t=self.file.add_paragraph(text)
        t.alignment=alignment
        return t

        
    def Save(self):
        self.file.save(self.loc)
        os.startfile(self.loc,"open")
        
    def SaveAsPDF(self):
        pass
 


class plot():
    """
Data : list (x and y), Series or Dataframe objects

kwargs:
    start_date
    
    end_date
    
    title=""
    
    ylabel=False, can be set to True to display labels associated with a Series object input or set to text directly
    
    xlabel=False, can be set to True to display labels associated with a Series object input or set to text directly
    
    width=1
    
    sharex=False, share x axis with previous plot
    
    sharey=False, share y axis with previous plot
    
    new_graph, default of True, set to False if you want the current Data to be
    plotted on the current subplot
    
    new_line, default = True, set to False if you want the current addition to
    be added horizontally on the report, ei for the graphs to be side by side
    instead of below the previous graph


https://www.python-course.eu/matplotlib_multiple_figures.php


@author: jlumpkin
"""
    def __init__(self,**kwargs):
        '''
        kwargs:
            title
            figsize
            debug_mode
        '''



        plt.rcParams.update({'figure.max_open_warning': 0})
        self._initialize_variables()
        
        self.title=kwargs.get("title",False)
        self.figsize=kwargs.get("figsize",None)#(8.5,11)

        self.debug_mode=kwargs.get("debug_mode",False)
      
        self.figure=None
        self.generated=False

    def generate_plot(self):
            plt.ioff()
            
            plt.style.use('seaborn')
            #plt.style.use('default')
            #plt.rcParams['figure.constrained_layout.use'] = True
            
            self.nrow=self.current_row+self._height_save
            
            prev_ax=None
            tolerance = 10 # points
            connected=False
            figure, ax=plt.subplots(figsize=self.figsize)
            
            self.figure=figure

            
            
            gridsize = ( self.nrow,self.ncol)
            plt.rcParams['axes.axisbelow'] = True
            for plot in self.plots:
                
                
                if plot["new_graph"]:
                    if  plot["sharex"]:
    
                        prev_ax=plt.subplot2grid(gridsize,(plot["Row"]-1,plot["Col"]-1),rowspan=plot["height"],colspan=plot["width"],sharex=prev_ax)
                        plt.setp(prev_ax.get_xticklabels(), visible=False)
                    elif plot["sharey"]:
    
                        prev_ax=plt.subplot2grid(gridsize,(plot["Row"]-1,plot["Col"]-1),rowspan=plot["height"],colspan=plot["width"],sharey=prev_ax)
                        plt.setp(prev_ax.get_yticklabels(), visible=False)
                    else:
    
                        prev_ax=plt.subplot2grid(gridsize,(plot["Row"]-1,plot["Col"]-1),rowspan=plot["height"],colspan=plot["width"])
                    plt.grid(True, linewidth=1)

                if plot["title"]:
                    prev_ax.set_title(plot["title"])
                if plot["xlabel"]:
                    if isinstance(plot["xlabel"],str):
                        prev_ax.set_xlabel(plot["xlabel"])
                    else:
                        prev_ax.set_xlabel(plot["Data"].index.name)
                if plot["ylabel"]:
                    if isinstance(plot["ylabel"],str):
                        prev_ax.set_ylabel(plot["ylabel"])
                    else:
                        prev_ax.set_ylabel(plot["Data"].index.name)
                
                #pdb.set_trace()
                if plot["Type"]==ePlot_type.line:
    
                    prev_ax.plot(plot["Data"], picker=tolerance,label=plot["name"],**plot["kwargs"])
                elif plot["Type"]==ePlot_type.scatter:
    
                    prev_ax.scatter(plot["Data"].index,plot["Data"].values, picker=tolerance,label=plot["name"],**plot["kwargs"])         
                elif plot["Type"]==ePlot_type.histogram:
                    prev_ax.hist(plot["Data"],bins=plot["bin"], picker=tolerance,label=plot["name"],**plot["kwargs"])     
                elif plot["Type"]==ePlot_type.span:
                    for span in plot["Data"]:
                        prev_ax.axvspan(span[0],span[1],alpha=.5,**plot["kwargs"])
                elif plot["Type"]==ePlot_type.bar:
                    prev_ax.bar(plot["Data"].index,plot["Data"].values,**plot["kwargs"])


                
                if plot["show_legend"]:
                    plt.legend(loc='best')
                if plot["func"] and not connected:
                    figure.func=plot["func"]
                    
                    connected=True

                if self.debug_mode:
                    prev_ax.set_facecolor("r")
                    prev_ax.axis("on")
               
                #plt.tight_layout()
                
                #plt.subplots_adjust(hspace=1)

            if self.title:
                figure.suptitle(self.title)

            self._initialize_variables()

            self.generated=True

    def plot_dict(self,**kwargs):
        plot={"Data":kwargs.get("Data",None),
            "new_graph":kwargs.get("new_graph",True),
            "Type":kwargs.get("Type",None),
            "Row":self.current_row,
            "Col":self.current_col,
            "title":kwargs.get("title",False),
            "name":kwargs.get("name",None),#if you pass in a list or ndarray, you can specify a name which will show up in a lengend
            "width":kwargs.get("width",1),
            "height":kwargs.get("height",1),
            "sharex":kwargs.get("sharex",False),
            "sharey":kwargs.get("sharey",False),
            "xlabel":kwargs.get("xlabel",None),
            "ylabel":kwargs.get("ylabel",None),
            "show_legend":kwargs.get("show_legend",False),
            "func":kwargs.get("func",None)}
        
        keys=plot.keys()
        for key in keys:#any kwarg that is not above get repackaged as kwargs in the plot_dict
            if key in kwargs:del kwargs[key]
        plot["kwargs"]=kwargs
            
        
        
        return plot


    def add_spans(self,Data,**kwargs):
        #Data must be a list of tuples with (start_date, end_date)

        self._adjust_plot_pos(**kwargs)

        plot=self.plot_dict(Data=Data,Type=ePlot_type.span,**kwargs)
        
        #plot["color"]=kwargs.get("color","green")
        
        if plot["height"]>self._height_save:
            self._height_save=plot["height"]

        self.plots.append(plot)


    
    def add_line_graph(self,*Data,**kwargs):


        Data=self._santize_data(*Data,**kwargs)

        self._adjust_plot_pos(**kwargs)

        plot=self.plot_dict(Data=Data,Type=ePlot_type.line,**kwargs)
        
        #plot["color"]=kwargs.get("color","red")
        
        if plot["height"]>self._height_save:
            self._height_save=plot["height"]

        self.plots.append(plot)


            
            
    def add_scatter_graph(self,*Data,**kwargs):

        Data=self._santize_data(*Data,**kwargs)

        self._adjust_plot_pos(**kwargs)
        
        plot=self.plot_dict(Data=Data,Type=ePlot_type.scatter,**kwargs)

        #plot["color"]=kwargs.get("color","red")
        #plot["s"]=kwargs.get("s",3)
        #plot["marker"]=kwargs.get("marker","o")#https://matplotlib.org/api/markers_api.html

        
        if plot["height"]>self._height_save:
            self._height_save=plot["height"]

        self.plots.append(plot)    
        
        
    def add_bar_graph(self,*Data,**kwargs):

        Data=self._santize_data(*Data,**kwargs)

        self._adjust_plot_pos(**kwargs)

        plot=self.plot_dict(Data=Data,Type=ePlot_type.bar,**kwargs)
        
        #plot["color"]=kwargs.get("color","red")
        
        if plot["height"]>self._height_save:
            self._height_save=plot["height"]

        self.plots.append(plot)


    def add_histogram_graph(self,*Data,**kwargs):

        Data=self._santize_data(*Data,**kwargs)

        self._adjust_plot_pos(**kwargs)

        

        kwargs["bin"]=kwargs.get("bin",6)#this keyword applies to histograms only
        plot=self.plot_dict(Data=Data,Type=ePlot_type.histogram,**kwargs)

        if plot["height"]>self._height_save:
            self._height_save=plot["height"]

        self.plots.append(plot) 


            
    def _initialize_variables(self):
        self.nrow, self.ncol, self.current_row, self.current_col =1,1,1,1
        self._height_save=1
        self.plots=list()#list of dicts
        
        self.first_plot=True
        
    def _adjust_plot_pos(self,**kwargs):
        if self.first_plot:
            self.first_plot=False
            self.nrow+=kwargs.get("height",1)-1
            return
        new_graph=kwargs.get("new_graph",True)
        
        if new_graph:

            new_line=kwargs.get("new_line",True)
            
            
            if new_line:#if this addition is stacking vertically
                self.current_col=1
                self.current_row+=self._height_save
                self.nrow+=self._height_save
                self._height_save=1
            else:#if this addition if stacking horizontally
                self.current_col+=1
                if self.current_col>self.ncol:
                    self.ncol+=1
                    

                
                
    def _santize_data(self,*Data,**kwargs):#insures data become a Series object

        if len(Data)==1  :
            if isinstance(Data[0],pd.Series) or isinstance(Data[0],pd.DataFrame):
                return Data[0]
            elif isinstance(Data[0],list):
                ret=pd.Series(data=Data[0],index=list(range(1,len(Data[0])+1)))
                ret.name=kwargs.get("name","")
                ret.index.name=""
                return ret

            else:
                raise Exception("Error, can only accept Series or Dataframe objects")

        elif len(Data)==2 :
            x=None
            y=None
            name=kwargs.get("name","")
            index_name=""
            if isinstance(Data[0],pd.Series):
                x=Data[0].values
                index_name=Data[0].name
                
            if isinstance(Data[1],pd.Series):
                y=Data[1].values
                name=Data[1].name
                
            if isinstance(Data[0],list):
                x=Data[0]
            elif isinstance(Data[0],range):
                x=list(Data[0])
                
                
                
            if isinstance(Data[1],list):
                y=Data[1]
            elif isinstance(Data[1],range):
                x=list(Data[1])                
            
            if (isinstance(Data[0],pd.Series) or isinstance(Data[0],list)or isinstance(Data[0],range)) and (isinstance(Data[1],pd.Series) or isinstance(Data[1],list)or isinstance(Data[1],range)):
                ret=pd.Series(data=y,index=x)
                ret.name=name
                ret.index.name=index_name
                return ret
            else:
                raise Exception("Error, when two arguments are passed in, they must both be Series, list, or range objects")
        else:
            raise Exception("Error, can only accept 1 or 2 arguments")
                    


    def _get_index(self,row,col,width):   
        if width==1:
        
            return (row-1)*self.ncol+col
        else:
            return((row-1)*self.ncol+col,(row-1)*self.ncol+col+width-1)





